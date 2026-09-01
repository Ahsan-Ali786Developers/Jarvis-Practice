from docx import Document as d
from datetime import datetime as dt

def create_receipt(stampValue, ChargedPrice, GovtFee, Profit):
	date = dt.now()
	doc = d()
	doc.add_heading("Stamp Receipt", level=1)
	doc.add_paragraph(f"Date: {date.strftime('%Y-%m-%d_%H:%M:%S')} ")
	doc.add_paragraph(f"Stamp Value: {stampValue} PKR")
	doc.add_paragraph(f"Amount Charged: {ChargedPrice} PKR")
	doc.save(f"receipt_{date.strftime('%Y-%m-%d_%H:%M:%S')}.docx")